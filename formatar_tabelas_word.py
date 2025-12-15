import os
import sys
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE_PATH = os.path.join("media", "word", "template tabela.docx")
STYLE_NAMES = ["Tabela de Grade 4 - Ênfase 1", "Grid Table 4 - Accent 1"]


def aplicar_estilo_sem_espacamento(paragraph):
    try:
        paragraph.style = 'No Spacing'
    except:
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1.0


def ajustar_automaticamente_tabela(table):
    try:
        for style_name in STYLE_NAMES:
            try:
                table.style = style_name
                break
            except Exception:
                continue
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    aplicar_estilo_sem_espacamento(paragraph)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = None
    except Exception as e:
        print(f"Erro ao ajustar tabela: {e}")


def remover_tabelas_existentes(doc):
    body = doc._element.body
    for tbl in doc.tables:
        body.remove(tbl._element)


def extrair_dados_tabela(orig_table):
    """
    Extrai os dados da tabela como lista de listas (linhas), incluindo cabeçalho.
    """
    dados = []
    for row in orig_table.rows:
        linha = []
        for cell in row.cells:
            # Junta todos os parágrafos da célula
            texto = "\n".join([p.text for p in cell.paragraphs]).strip()
            linha.append(texto)
        dados.append(linha)
    return dados


def ordenar_dados_tabela(dados):
    """
    Ordena os dados da tabela pela segunda coluna (data, índice 1) e depois pela primeira (placa, índice 0).
    Considera a primeira linha como cabeçalho.
    """
    if len(dados) <= 1:
        return dados
    cabecalho = dados[0]
    linhas = dados[1:]
    idx_data = 1  # Segunda coluna
    idx_placa = 0  # Primeira coluna
    def parse_data(s):
        for fmt in ("%d-%b-%y", "%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%d/%m/%y", "%d-%b-%Y"):
            try:
                return datetime.strptime(s, fmt)
            except Exception:
                continue
        return None
    def sort_key(x):
        data = parse_data(x[idx_data])
        placa = x[idx_placa]
        return (data if data is not None else datetime.max, placa)
    linhas_ordenadas = sorted(linhas, key=sort_key)
    return [cabecalho] + linhas_ordenadas


def set_borda_simples(table):
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def copiar_tabela_ordenada(orig_table, doc_destino):
    dados = extrair_dados_tabela(orig_table)
    print('Dados extraídos da tabela:')
    for linha in dados:
        print(linha)
    dados_ordenados = ordenar_dados_tabela(dados)
    rows = len(dados_ordenados)
    cols = max(len(linha) for linha in dados_ordenados) if rows > 0 else 0
    style_aplicado = None
    for style_name in STYLE_NAMES:
        try:
            new_table = doc_destino.add_table(rows=0, cols=cols, style=style_name)
            style_aplicado = style_name
            break
        except Exception:
            continue
    if not style_aplicado:
        new_table = doc_destino.add_table(rows=0, cols=cols)
    for linha in dados_ordenados:
        row_cells = new_table.add_row().cells
        for j, valor in enumerate(linha):
            row_cells[j].text = valor if valor is not None else ""
    set_borda_simples(new_table)
    print(f'Tabela criada: {len(new_table.rows)} linhas x {len(new_table.columns)} colunas')
    return new_table


def formatar_tabelas_arquivo(caminho_arquivo):
    try:
        doc_origem = Document(caminho_arquivo)
        doc_destino = Document(TEMPLATE_PATH)
        remover_tabelas_existentes(doc_destino)
        num_tabelas = len(doc_origem.tables)
        print(f"Encontradas {num_tabelas} tabelas no documento.")
        if num_tabelas == 0:
            print("Nenhuma tabela encontrada no documento.")
            return False
        for i, table in enumerate(doc_origem.tables, 1):
            print(f"Copiando, ordenando e formatando tabela {i}/{num_tabelas}...")
            new_table = copiar_tabela_ordenada(table, doc_destino)
            ajustar_automaticamente_tabela(new_table)
            doc_destino.add_paragraph("")  # Espaço entre tabelas
        nome_arquivo = os.path.basename(caminho_arquivo)
        nome_base, extensao = os.path.splitext(nome_arquivo)
        caminho_saida = os.path.join(os.path.dirname(caminho_arquivo), f"{nome_base}_formatado{extensao}")
        doc_destino.save(caminho_saida)
        print(f"Documento salvo como: {caminho_saida}")
        return True
    except Exception as e:
        print(f"Erro ao processar o arquivo: {e}")
        return False


def selecionar_arquivo_word():
    root = tk.Tk()
    root.withdraw()
    arquivo = filedialog.askopenfilename(
        title="Selecione um arquivo Word (.docx)",
        filetypes=[("Arquivos Word", "*.docx"), ("Todos os arquivos", "*.*")]
    )
    return arquivo


def main():
    print("=== Formatador de Tabelas Word ===")
    print("Este script formata tabelas com estilo 'Sem Espaçamento', centralização, ajuste automático e ordenação.")
    print()
    if len(sys.argv) > 1:
        caminho_arquivo = sys.argv[1]
        if not os.path.exists(caminho_arquivo):
            print(f"Arquivo não encontrado: {caminho_arquivo}")
            return
    else:
        print("Selecione o arquivo Word que deseja formatar...")
        caminho_arquivo = selecionar_arquivo_word()
        if not caminho_arquivo:
            print("Nenhum arquivo selecionado.")
            return
    print(f"Processando arquivo: {caminho_arquivo}")
    sucesso = formatar_tabelas_arquivo(caminho_arquivo)
    if sucesso:
        print("\n✅ Formatação concluída com sucesso!")
        messagebox.showinfo("Sucesso", "Tabelas formatadas com sucesso!")
    else:
        print("\n❌ Erro na formatação.")
        messagebox.showerror("Erro", "Erro ao formatar as tabelas.")

if __name__ == "__main__":
    main() 