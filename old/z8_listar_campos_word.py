import os
import re
import sqlite3
from docx import Document

# Caminho do arquivo e do banco de dados
template_path = 'media/template/Template Relatório.docx'
db_path = 'banco_dados.db'

# Conectar ao banco de dados e criar cursor
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

# Verificar se a tabela já existe, caso contrário, criar
table_name = 'relatorio_slu'
cursor.execute(f'''
    CREATE TABLE IF NOT EXISTS {table_name} (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        palavra TEXT UNIQUE,
        tipo_informacao TEXT,
        pagina INTEGER,
        origem TEXT
    )
''')
print("Tabela verificada/criada com sucesso.")

# Abrir o documento
if not os.path.exists(template_path):
    print("Erro: o arquivo Template Relatório.docx não foi encontrado.")
else:
    document = Document(template_path)
    padrao = r'##\w+##'  # Regex para encontrar palavras entre ##

    registros_encontrados = 0
    registros_duplicados = 0
    erros_ignorados = 0

    def extrair_texto_paragrafos(elemento):
        """Extrai o texto de todos os parágrafos de um elemento."""
        texto = []
        for paragrafo in elemento.paragraphs:
            texto.append(paragrafo.text)
        return '\n'.join(texto)

    # Processar parágrafos e caixas de texto
    for paragrafo in document.paragraphs:
        palavras_encontradas = re.findall(padrao, paragrafo.text)
        for palavra in palavras_encontradas:
            tipo_informacao = "Desconhecido"
            origem = "Texto Principal"
            try:
                cursor.execute(f'''
                    INSERT INTO {table_name} (palavra, tipo_informacao, pagina, origem)
                    VALUES (?, ?, ?, ?)
                ''', (palavra, tipo_informacao, 1, origem))
                registros_encontrados += 1
                print(f"Registro adicionado: {palavra} na origem {origem}")
            except sqlite3.IntegrityError:
                registros_duplicados += 1
                print(f"Registro duplicado ignorado: {palavra}")
            except Exception as e:
                erros_ignorados += 1
                print(f"Erro ao inserir {palavra}: {e}")

    # Processar possíveis caixas de texto
    for shape in document.inline_shapes:
        if shape.text_frame:
            texto_caixa = shape.text_frame.text
            palavras_encontradas = re.findall(padrao, texto_caixa)
            for palavra in palavras_encontradas:
                tipo_informacao = "Desconhecido"
                origem = "Caixa de Texto"
                try:
                    cursor.execute(f'''
                        INSERT INTO {table_name} (palavra, tipo_informacao, pagina, origem)
                        VALUES (?, ?, ?, ?)
                    ''', (palavra, tipo_informacao, 1, origem))
                    registros_encontrados += 1
                    print(f"Registro adicionado: {palavra} na origem {origem}")
                except sqlite3.IntegrityError:
                    registros_duplicados += 1
                    print(f"Registro duplicado ignorado: {palavra}")
                except Exception as e:
                    erros_ignorados += 1
                    print(f"Erro ao inserir {palavra}: {e}")

    # Salvar as alterações no banco de dados
    conn.commit()
    print(f"\nProcesso concluído. Registros adicionados: {registros_encontrados}, Duplicados ignorados: {registros_duplicados}, Erros ignorados: {erros_ignorados}")

# Fechar a conexão com o banco de dados
conn.close()
