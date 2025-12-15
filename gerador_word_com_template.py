import os
import sqlite3
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from dateutil.relativedelta import relativedelta

# --- CONFIGURAÇÕES ---
CAMINHO_BD = "banco_dados_completo/banco_dados.db"
TABELA_BD = "placas_completas_slu_bh"
TEMPLATE_PATH = "relatorios/template_tabela.docx"  # O SEU TEMPLATE COM A TABELA MODELO
OUTPUT_PATH = "relatorios/relatorio_final.docx"

# Dicionário de mapeamento (copiado do seu script)
placa_to_table = {
    "PR01": "Recalques Celula Pesquisa", "PR02": "Recalques Celula Pesquisa",
    "PR03": "Recalques Celula Pesquisa", "PR04": "Recalques Celula Pesquisa",
    "PR05": "Recalques Celula Pesquisa", "PR06": "Recalques Celula Pesquisa",
    "PR07": "Recalques Celula Pesquisa", "PR08": "Recalques Celula Pesquisa",
    "PR09": "Recalques Celula Pesquisa", "PR10": "Recalques Celula Pesquisa",
    "PR11": "Recalques Celula Pesquisa", "PR12": "Recalques Celula Pesquisa",
    "PR13": "Recalques Celula Pesquisa", "PR14": "Recalques Celula Pesquisa",
    "PR15": "Recalques Celula Pesquisa",
    "PR 1.1": "Placa de Recalque AC 01", "PR 1.2": "Placa de Recalque AC 01",
    "PR 1.3": "Placa de Recalque AC 01", "PR 1.4": "Placa de Recalque AC 01",
    "PR 1.5": "Placa de Recalque AC 01", "D1": "Placa de Recalque AC 01",
    "D2": "Placa de Recalque AC 01", "PR 1.21": "Placa de Recalque AC 01",
    "PR 1.22": "Placa de Recalque AC 01", "PR 1.23": "Placa de Recalque AC 01",
    "PR 1.24": "Placa de Recalque AC 01", "PR 1.25": "Placa de Recalque AC 01",
    "PR 1.31": "Placa de Recalque AC 01", "PR 1.32": "Placa de Recalque AC 01",
    "PR 1.33": "Placa de Recalque AC 01", "PR 1.34": "Placa de Recalque AC 01",
    "PR 1.35": "Placa de Recalque AC 01", "PR 1.41": "Placa de Recalque AC 01",
    "PR 1.42": "Placa de Recalque AC 01", "PR 1.43": "Placa de Recalque AC 01",
    "PR 1.44": "Placa de Recalque AC 01", "PR 1.45": "Placa de Recalque AC 01",
    "PR 3.21": "Placa de Recalque AC 03", "PR 3.22": "Placa de Recalque AC 03",
    "PR 3.23": "Placa de Recalque AC 03", "PR 3.24": "Placa de Recalque AC 03",
    "PR 3.25": "Placa de Recalque AC 03", "PR 3.26": "Placa de Recalque AC 03",
    "PR 3.27": "Placa de Recalque AC 03", "PR 3.28": "Placa de Recalque AC 03",
    "PR 3.29": "Placa de Recalque AC 03", "PR 3.30": "Placa de Recalque AC 03",
    "PR 3.31": "Placa de Recalque AC 03",
    "PR 4.1": "Placa de Recalque AC 04", "PR 4.2": "Placa de Recalque AC 04",
    "PR 4.3": "Placa de Recalque AC 04", "PR 4.4": "Placa de Recalque AC 04",
    "PR 4.5": "Placa de Recalque AC 04", "PR 4.6": "Placa de Recalque AC 04",
    "PR 4.7": "Placa de Recalque AC 04", "PR 4.8": "Placa de Recalque AC 04",
    "PR 4.9": "Placa de Recalque AC 04", "PR 4.10": "Placa de Recalque AC 04",
    "PR 4.11": "Placa de Recalque AC 04", "PR 4.21": "Placa de Recalque AC 04",
    "PR 4.22": "Placa de Recalque AC 04", "PR 4.23": "Placa de Recalque AC 04",
    "PR 4.24": "Placa de Recalque AC 04", "PR 4.25": "Placa de Recalque AC 04",
    "PR 4.26": "Placa de Recalque AC 04", "PR 4.27": "Placa de Recalque AC 04",
    "PR 4.28": "Placa de Recalque AC 04", "PR 4.29": "Placa de Recalque AC 04",
    "PR 4.30": "Placa de Recalque AC 04", "PR 4.31": "Placa de Recalque AC 04",
    "PR 4.41": "Placa de Recalque AC 04", "PR 4.42": "Placa de Recalque AC 04",
    "PR 4.43": "Placa de Recalque AC 04", "PR 4.44": "Placa de Recalque AC 04",
    "PR 4.45": "Placa de Recalque AC 04", "PR 4.46": "Placa de Recalque AC 04",
    "PR 4.47": "Placa de Recalque AC 04", "PR 4.48": "Placa de Recalque AC 04",
    "PR 4.49": "Placa de Recalque AC 04", "PR 4.50": "Placa de Recalque AC 04",
    "PR 4.51": "Placa de Recalque AC 04",
    "PR 5.1": "Placa de Recalque AC 05", "PR 5.2": "Placa de Recalque AC 05",
    "PR 5.3": "Placa de Recalque AC 05", "PR 5.4": "Placa de Recalque AC 05",
    "PR 5.5": "Placa de Recalque AC 05", "PR 5.6": "Placa de Recalque AC 05",
    "PR 5.7": "Placa de Recalque AC 05", "PR 5.8": "Placa de Recalque AC 05",
    "PR 5.9": "Placa de Recalque AC 05", "PR 5.10": "Placa de Recalque AC 05",
    "PR 5.11": "Placa de Recalque AC 05", "PR 5.21": "Placa de Recalque AC 05",
    "PR 5.22": "Placa de Recalque AC 05", "PR 5.23": "Placa de Recalque AC 05",
    "PR 5.24": "Placa de Recalque AC 05", "PR 5.25": "Placa de Recalque AC 05",
    "PR 5.26": "Placa de Recalque AC 05", "PR 5.27": "Placa de Recalque AC 05",
    "PR 5.28": "Placa de Recalque AC 05", "PR 5.29": "Placa de Recalque AC 05",
    "PR 5.30": "Placa de Recalque AC 05", "PR 5.41": "Placa de Recalque AC 05",
    "PR 5.42": "Placa de Recalque AC 05", "PR 5.43": "Placa de Recalque AC 05",
    "PR 5.44": "Placa de Recalque AC 05", "PR 5.45": "Placa de Recalque AC 05",
    "PR 5.46": "Placa de Recalque AC 05", "PR 5.47": "Placa de Recalque AC 05",
    "PR 5.48": "Placa de Recalque AC 05", "PR 5.52": "Placa de Recalque AC 05",
    "PR 5.53": "Placa de Recalque AC 05", "PR 5.54": "Placa de Recalque AC 05",
    "PR 5.55": "Placa de Recalque AC 05", "PR 5.56": "Placa de Recalque AC 05",
    "PR A.1": "Placa de Recalque AMPLIAÇÃO", "PR A.2": "Placa de Recalque AMPLIAÇÃO",
    "PR A.3": "Placa de Recalque AMPLIAÇÃO", "PR A.4": "Placa de Recalque AMPLIAÇÃO",
    "PR A.5": "Placa de Recalque AMPLIAÇÃO", "PR A.6": "Placa de Recalque AMPLIAÇÃO",
    "PR A.7": "Placa de Recalque AMPLIAÇÃO", "PR A.8": "Placa de Recalque AMPLIAÇÃO",
    "PR A.9": "Placa de Recalque AMPLIAÇÃO",
    "PR 1.1": "Placa de Recalque Emergencial", "PR 1.2": "Placa de Recalque Emergencial",
    "PR 1.3": "Placa de Recalque Emergencial", "PR 1.4": "Placa de Recalque Emergencial",
    "PR 1.5": "Placa de Recalque Emergencial", "PR 2.1": "Placa de Recalque Emergencial",
    "PR 2.2": "Placa de Recalque Emergencial", "PR 2.3": "Placa de Recalque Emergencial",
    "PR 2.4": "Placa de Recalque Emergencial", "PR 2.5": "Placa de Recalque Emergencial",
    "PR 3.1": "Placa de Recalque Emergencial", "PR 3.2": "Placa de Recalque Emergencial",
    "PR 3.3": "Placa de Recalque Emergencial", "PR 3.4": "Placa de Recalque Emergencial",
}

# Ordem das tabelas no relatório final
tabelas_relatorio = [
    "Recalques Celula Pesquisa", "Placa de Recalque Emergencial",
    "Placa de Recalque AC 01", "Placa de Recalque AC 03",
    "Placa de Recalque AC 04", "Placa de Recalque AC 05",
    "Placa de Recalque AMPLIAÇÃO",
]

# --- FUNÇÕES AUXILIARES ---

def clone_table(table, doc):
    """Cria uma nova tabela no doc com o mesmo número de colunas, copia o cabeçalho, aplica o estilo 'Table Grid', centraliza tudo e deixa o cabeçalho em negrito."""
    n_cols = len(table.columns)
    new_table = doc.add_table(rows=1, cols=n_cols)
    # Copia o cabeçalho
    for i, cell in enumerate(table.rows[0].cells):
        new_cell = new_table.rows[0].cells[i]
        new_cell.text = cell.text
        # Centralizar e negritar o cabeçalho
        for paragraph in new_cell.paragraphs:
            paragraph.alignment = 1  # 1 = center
            for run in paragraph.runs:
                run.bold = True
    # Aplica o estilo de tabela do Word
    new_table.style = 'Table Grid'
    return new_table


def formatar_data_pt_br_iso(yyyy_mm_dd: str) -> str:
    if not yyyy_mm_dd or len(str(yyyy_mm_dd)) < 10: return str(yyyy_mm_dd)
    try:
        dt = datetime.strptime(str(yyyy_mm_dd)[:10], '%Y-%m-%d')
        meses_pt = ["jan", "fev", "mar", "abr", "mai", "jun", "jul", "ago", "set", "out", "nov", "dez"]
        return f"{dt.day}-{meses_pt[dt.month-1]}-{dt.strftime('%y')}"
    except:
        return str(yyyy_mm_dd)

def formatar_cota(valor: float) -> str:
    if valor is None: return ""
    return f"{valor:.3f}".replace('.', ',')

def formatar_coordenada(valor: float) -> str:
    if valor is None: return ""
    return f"{valor:,.3f}".replace(",", "X").replace(".", ",").replace("X", ".")

# --- LÓGICA PRINCIPAL ---

def gerar_relatorio():
    print("--- Iniciando Gerador de Relatório ---")

    # 1. Verificar se o template existe
    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERRO: Arquivo template '{TEMPLATE_PATH}' não encontrado.")
        print("Por favor, crie este arquivo com uma tabela modelo formatada.")
        return

    # 2. Carregar dados do banco
    print(f"Conectando ao banco de dados: {CAMINHO_BD}")
    conn = sqlite3.connect(CAMINHO_BD)
    df = pd.read_sql_query(f"SELECT * FROM {TABELA_BD}", conn)
    conn.close()
    print("Dados carregados com sucesso.")

    # 3. Processar e filtrar dados
    df['data_datetime'] = pd.to_datetime(df['data'], errors='coerce', format='%Y-%m-%d')
    data_maxima = df['data_datetime'].max()
    if pd.isnull(data_maxima):
        print("ERRO: Nenhuma data válida encontrada no banco de dados.")
        return
    
    cinco_meses_atras = data_maxima - relativedelta(months=5)
    df_filtrado = df[df['data_datetime'] >= cinco_meses_atras].copy()
    df_filtrado['nome_tabela_final'] = df_filtrado['placa'].map(placa_to_table)
    print(f"Dados filtrados para o período a partir de {cinco_meses_atras.strftime('%d/%m/%Y')}.")

    # 4. Preparar documento Word a partir do template
    doc = Document() # Documento final em branco
    template_doc = Document(TEMPLATE_PATH)
    if not template_doc.tables:
        print(f"ERRO: Nenhuma tabela encontrada no template '{TEMPLATE_PATH}'.")
        return
    template_table = template_doc.tables[0]
    print("Template do Word carregado.")

    # 5. Gerar tabelas
    tabelas_geradas = []
    for nome_tabela in tabelas_relatorio:
        print(f"\nProcessando tabela: '{nome_tabela}'")
        df_tab = df_filtrado[df_filtrado['nome_tabela_final'] == nome_tabela].copy()

        if df_tab.empty:
            print(" -> Tabela vazia, pulando.")
            tabelas_geradas.append(None)
            continue

        # ** ORDENAÇÃO CORRIGIDA PELA DATA **
        df_tab.sort_values(by=["data_datetime"], inplace=True)
        print(f" -> Encontrados {len(df_tab)} registros. Ordenando por data.")

        # Adicionar título
        doc.add_paragraph(nome_tabela.upper(), style='Heading 1')
        
        # Clonar a tabela do template
        new_table = clone_table(template_table, doc)
        
        # Limpar linhas de exemplo (se houver) e manter cabeçalho
        while len(new_table.rows) > 1:
            new_table._tbl.remove(new_table.rows[-1]._tr)

        # Popular a tabela com dados
        for _, row in df_tab.iterrows():
            row_cells = new_table.add_row().cells
            row_cells[0].text = str(row['placa'] or "")
            row_cells[1].text = formatar_data_pt_br_iso(row['data'])
            row_cells[2].text = formatar_cota(row['elevacao'])
            row_cells[3].text = formatar_coordenada(row['coordenada_este'])
            row_cells[4].text = formatar_coordenada(row['coordenada_norte'])
            # Centralizar texto das células de dados
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # 1 = center
        # Remover bloco de cópia de shading do cabeçalho
        tabelas_geradas.append(new_table)
        doc.add_page_break()
        print(" -> Tabela adicionada ao documento.")

    # Nova abordagem: criar novo documento e inserir tabelas nos marcadores
    print("\n--- Gerando documento final com inserção robusta das tabelas ---")
    from docx.shared import RGBColor, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    relatorio_template_path = "media/word/template_relatorio.docx"
    relatorio_saida_path = "relatorios/relatorio_final_completo.docx"
    doc_template = Document(relatorio_template_path)
    doc_final = Document()
    marcador_base = "Tabela_"
    tabela_idx = 0
    i = 0
    while i < len(doc_template.paragraphs):
        p = doc_template.paragraphs[i]
        texto = p.text.replace('“','').replace('”','').replace('"','').strip()
        marcador = f"Tabela_{tabela_idx+1}"
        if marcador == texto or marcador in texto:
            # Inserir tabela formatada
            tabela_dados = tabelas_geradas[tabela_idx]
            if tabela_dados is not None:
                tbl_data = []
                for row in tabela_dados.rows:
                    linha = []
                    for cell in row.cells:
                        linha.append(cell.text)
                    tbl_data.append(linha)
                rows = len(tbl_data)
                cols = len(tbl_data[0]) if rows > 0 else 0
                table = doc_final.add_table(rows=0, cols=cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # Cabeçalho
                hdr_cells = table.add_row().cells
                for j, valor in enumerate(tbl_data[0]):
                    paragraph = hdr_cells[j].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(valor)
                    run.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)
                    run.font.size = Pt(11)
                    tcPr = hdr_cells[j]._tc.get_or_add_tcPr()
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), '4472C4')
                    tcPr.append(shd)
                # Dados
                for linha in tbl_data[1:]:
                    row_cells = table.add_row().cells
                    for j, valor in enumerate(linha):
                        paragraph = row_cells[j].paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(valor)
                        run.font.size = Pt(11)
                doc_final.add_paragraph("")
            tabela_idx += 1
            i += 1
            continue
        # Copiar parágrafo normalmente
        p_final = doc_final.add_paragraph()
        for run in p.runs:
            r = p_final.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.size = run.font.size
            r.font.color.rgb = run.font.color.rgb if run.font.color else None
        p_final.alignment = p.alignment
        i += 1
    doc_final.save(relatorio_saida_path)
    print(f"Relatório final salvo em: {relatorio_saida_path}")

    # 6. Salvar o documento final
    try:
        doc.save(OUTPUT_PATH)
        print(f"\n--- SUCESSO! ---")
        print(f"Relatório salvo como: '{OUTPUT_PATH}'")
    except Exception as e:
        print(f"\n--- ERRO AO SALVAR ---")
        print(f"Não foi possível salvar o arquivo: {e}")

if __name__ == "__main__":
    gerar_relatorio() 