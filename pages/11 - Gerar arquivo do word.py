import streamlit as st
import sqlite3
import pandas as pd
import os
from datetime import datetime
from dateutil.relativedelta import relativedelta

# Bibliotecas para escrever no Word (python-docx)
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn

# ----------------------------------------------------------------
# Ajuste este caminho conforme seu diretório de arquivos Word:
DIRETORIO_ARQUIVOS = "media/word"

# Ajuste o caminho para seu banco de dados:
CAMINHO_BD = "banco_dados_completo/banco_dados.db"
TABELA_BD = "placas_completas_slu_bh"

# Dicionário que mapeia cada placa ao nome da Tabela que ela pertence
# (preencha TODO o mapeamento conforme sua tabela de referência)
placa_to_table = {
    # Recalques Celula Pesquisa
    "PR01": "Recalques Celula Pesquisa",
    "PR02": "Recalques Celula Pesquisa",
    "PR03": "Recalques Celula Pesquisa",
    "PR04": "Recalques Celula Pesquisa",
    "PR05": "Recalques Celula Pesquisa",
    "PR06": "Recalques Celula Pesquisa",
    "PR07": "Recalques Celula Pesquisa",
    "PR08": "Recalques Celula Pesquisa",
    "PR09": "Recalques Celula Pesquisa",
    "PR10": "Recalques Celula Pesquisa",
    "PR11": "Recalques Celula Pesquisa",
    "PR12": "Recalques Celula Pesquisa",
    "PR13": "Recalques Celula Pesquisa",
    "PR14": "Recalques Celula Pesquisa",
    "PR15": "Recalques Celula Pesquisa",

    # Placa de Recalque AC 01
    "PR 1.1": "Placa de Recalque AC 01",
    "PR 1.2": "Placa de Recalque AC 01",
    "PR 1.3": "Placa de Recalque AC 01",
    "PR 1.4": "Placa de Recalque AC 01",
    "PR 1.5": "Placa de Recalque AC 01",
    "D1": "Placa de Recalque AC 01",
    "D2": "Placa de Recalque AC 01",
    "PR 1.21": "Placa de Recalque AC 01",
    "PR 1.22": "Placa de Recalque AC 01",
    "PR 1.23": "Placa de Recalque AC 01",
    "PR 1.24": "Placa de Recalque AC 01",
    "PR 1.25": "Placa de Recalque AC 01",
    "PR 1.31": "Placa de Recalque AC 01",
    "PR 1.32": "Placa de Recalque AC 01",
    "PR 1.33": "Placa de Recalque AC 01",
    "PR 1.34": "Placa de Recalque AC 01",
    "PR 1.35": "Placa de Recalque AC 01",
    "PR 1.41": "Placa de Recalque AC 01",
    "PR 1.42": "Placa de Recalque AC 01",
    "PR 1.43": "Placa de Recalque AC 01",
    "PR 1.44": "Placa de Recalque AC 01",
    "PR 1.45": "Placa de Recalque AC 01",

    # Placa de Recalque AC 03
    "PR 3.21": "Placa de Recalque AC 03",
    "PR 3.22": "Placa de Recalque AC 03",
    "PR 3.23": "Placa de Recalque AC 03",
    "PR 3.24": "Placa de Recalque AC 03",
    "PR 3.25": "Placa de Recalque AC 03",
    "PR 3.26": "Placa de Recalque AC 03",
    "PR 3.27": "Placa de Recalque AC 03",
    "PR 3.28": "Placa de Recalque AC 03",
    "PR 3.29": "Placa de Recalque AC 03",
    "PR 3.30": "Placa de Recalque AC 03",
    "PR 3.31": "Placa de Recalque AC 03",

    # Placa de Recalque AC 04
    "PR 4.1": "Placa de Recalque AC 04",
    "PR 4.2": "Placa de Recalque AC 04",
    "PR 4.3": "Placa de Recalque AC 04",
    "PR 4.4": "Placa de Recalque AC 04",
    "PR 4.5": "Placa de Recalque AC 04",
    "PR 4.6": "Placa de Recalque AC 04",
    "PR 4.7": "Placa de Recalque AC 04",
    "PR 4.8": "Placa de Recalque AC 04",
    "PR 4.9": "Placa de Recalque AC 04",
    "PR 4.10": "Placa de Recalque AC 04",
    "PR 4.11": "Placa de Recalque AC 04",
    "PR 4.21": "Placa de Recalque AC 04",
    "PR 4.22": "Placa de Recalque AC 04",
    "PR 4.23": "Placa de Recalque AC 04",
    "PR 4.24": "Placa de Recalque AC 04",
    "PR 4.25": "Placa de Recalque AC 04",
    "PR 4.26": "Placa de Recalque AC 04",
    "PR 4.27": "Placa de Recalque AC 04",
    "PR 4.28": "Placa de Recalque AC 04",
    "PR 4.29": "Placa de Recalque AC 04",
    "PR 4.30": "Placa de Recalque AC 04",
    "PR 4.31": "Placa de Recalque AC 04",
    "PR 4.41": "Placa de Recalque AC 04",
    "PR 4.42": "Placa de Recalque AC 04",
    "PR 4.43": "Placa de Recalque AC 04",
    "PR 4.44": "Placa de Recalque AC 04",
    "PR 4.45": "Placa de Recalque AC 04",
    "PR 4.46": "Placa de Recalque AC 04",
    "PR 4.47": "Placa de Recalque AC 04",
    "PR 4.48": "Placa de Recalque AC 04",
    "PR 4.49": "Placa de Recalque AC 04",
    "PR 4.50": "Placa de Recalque AC 04",
    "PR 4.51": "Placa de Recalque AC 04",

    # Placa de Recalque AC 05
    "PR 5.1": "Placa de Recalque AC 05",
    "PR 5.2": "Placa de Recalque AC 05",
    "PR 5.3": "Placa de Recalque AC 05",
    "PR 5.4": "Placa de Recalque AC 05",
    "PR 5.5": "Placa de Recalque AC 05",
    "PR 5.6": "Placa de Recalque AC 05",
    "PR 5.7": "Placa de Recalque AC 05",
    "PR 5.8": "Placa de Recalque AC 05",
    "PR 5.9": "Placa de Recalque AC 05",
    "PR 5.10": "Placa de Recalque AC 05",
    "PR 5.11": "Placa de Recalque AC 05",
    "PR 5.21": "Placa de Recalque AC 05",
    "PR 5.22": "Placa de Recalque AC 05",
    "PR 5.23": "Placa de Recalque AC 05",
    "PR 5.24": "Placa de Recalque AC 05",
    "PR 5.25": "Placa de Recalque AC 05",
    "PR 5.26": "Placa de Recalque AC 05",
    "PR 5.27": "Placa de Recalque AC 05",
    "PR 5.28": "Placa de Recalque AC 05",
    "PR 5.29": "Placa de Recalque AC 05",
    "PR 5.30": "Placa de Recalque AC 05",
    "PR 5.41": "Placa de Recalque AC 05",
    "PR 5.42": "Placa de Recalque AC 05",
    "PR 5.43": "Placa de Recalque AC 05",
    "PR 5.44": "Placa de Recalque AC 05",
    "PR 5.45": "Placa de Recalque AC 05",
    "PR 5.46": "Placa de Recalque AC 05",
    "PR 5.47": "Placa de Recalque AC 05",
    "PR 5.48": "Placa de Recalque AC 05",
    "PR 5.52": "Placa de Recalque AC 05",
    "PR 5.53": "Placa de Recalque AC 05",
    "PR 5.54": "Placa de Recalque AC 05",
    "PR 5.55": "Placa de Recalque AC 05",
    "PR 5.56": "Placa de Recalque AC 05",

    # Placa de Recalque AMPLIAÇÃO
    "PR A.1": "Placa de Recalque AMPLIAÇÃO",
    "PR A.2": "Placa de Recalque AMPLIAÇÃO",
    "PR A.3": "Placa de Recalque AMPLIAÇÃO",
    "PR A.4": "Placa de Recalque AMPLIAÇÃO",
    "PR A.5": "Placa de Recalque AMPLIAÇÃO",
    "PR A.6": "Placa de Recalque AMPLIAÇÃO",
    "PR A.7": "Placa de Recalque AMPLIAÇÃO",
    "PR A.8": "Placa de Recalque AMPLIAÇÃO",
    "PR A.9": "Placa de Recalque AMPLIAÇÃO",

    # Placa de Recalque Emergencial
    "PR 1.1": "Placa de Recalque Emergencial",
    "PR 1.2": "Placa de Recalque Emergencial",
    "PR 1.3": "Placa de Recalque Emergencial",
    "PR 1.4": "Placa de Recalque Emergencial",
    "PR 1.5": "Placa de Recalque Emergencial",
    "PR 2.1": "Placa de Recalque Emergencial",
    "PR 2.2": "Placa de Recalque Emergencial",
    "PR 2.3": "Placa de Recalque Emergencial",
    "PR 2.4": "Placa de Recalque Emergencial",
    "PR 2.5": "Placa de Recalque Emergencial",
    "PR 3.1": "Placa de Recalque Emergencial",
    "PR 3.2": "Placa de Recalque Emergencial",
    "PR 3.3": "Placa de Recalque Emergencial",
    "PR 3.4": "Placa de Recalque Emergencial",
}

del placa_to_table["PR 3.3"]

# Lista (ou dicionário) com as 7 grandes Tabelas (na ordem desejada):
tabelas_relatorio = [
    "Recalques Celula Pesquisa",
    "Placa de Recalque Emergencial",
    "Placa de Recalque AC 01",
    "Placa de Recalque AC 03",
    "Placa de Recalque AC 04",
    "Placa de Recalque AC 05",
    "Placa de Recalque AMPLIAÇÃO",
]

# ----------------------------------------------------------------
# Funções auxiliares de formatação:

def formatar_data_pt_br_iso(yyyy_mm_dd: str) -> str:
    """
    Converte 'YYYY-MM-DD' para 'dd-mês-yy' em pt-BR, ex.: '10-fev-25'.
    """
    if not yyyy_mm_dd or len(yyyy_mm_dd) < 10:
        return yyyy_mm_dd  # Retorna como está se não for formato esperado
    
    ano = yyyy_mm_dd[0:4]
    mes = yyyy_mm_dd[5:7]
    dia = yyyy_mm_dd[8:10]

    # Dicionário simples de mês em pt-BR (abreviado)
    meses_pt = {
        "01": "jan", "02": "fev", "03": "mar", "04": "abr",
        "05": "mai", "06": "jun", "07": "jul", "08": "ago",
        "09": "set", "10": "out", "11": "nov", "12": "dez"
    }

    mes_pt = meses_pt.get(mes, mes)
    # Pegamos apenas os dois últimos dígitos do ano
    ano_2d = ano[-2:]
    return f"{dia}-{mes_pt}-{ano_2d}"


def formatar_cota(valor: float) -> str:
    """
    Formato: 000,000 (3 casas decimais, separador decimal vírgula).
    Sem ponto de milhar para a cota.
    Ex: 934.051 -> "934,051"
    """
    if valor is None:
        return ""
    return f"{valor:0.3f}".replace('.', ',')


def formatar_coordenada(valor: float) -> str:
    """
    Formato: separar milhar com ponto e usar vírgula como decimal,
    sempre com 3 casas decimais. Exemplo:
      602756.6407 -> "602.756,641"
      7797363.0474 -> "7.797.363,047"
    """
    if valor is None:
        return ""
    # arredonda para 3 decimais
    v = round(valor, 3)
    # fatia parte inteira e decimal
    parte_inteira = int(v)
    parte_decimal = abs(v - parte_inteira)
    # formata a parte inteira com pontos (.) entre cada milhar
    parte_inteira_str = f"{abs(parte_inteira):,}".replace(",", ".")
    # formata a parte decimal após a vírgula
    parte_decimal_str = f"{parte_decimal:.3f}"[1:].replace(".", ",")  # vira ',xxx'
    
    # recoloca sinal, se negativo
    sinal = "-" if v < 0 else ""
    return f"{sinal}{parte_inteira_str}{parte_decimal_str}"


def set_repeat_table_header(row):
    """
    Força a linha de cabeçalho a repetir nas páginas seguintes do Word.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)
    return row

# Adicionar esta função para sanitizar nomes de arquivos
def sanitize_filename(filename):
    """Remove caracteres inválidos para nomes de arquivos no Windows."""
    import re
    # Substitui caracteres inválidos por underscore
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
    # Remove acentos
    import unicodedata
    sanitized = unicodedata.normalize('NFKD', sanitized).encode('ASCII', 'ignore').decode('ASCII')
    return sanitized

# ----------------------------------------------------------------
# Layout do Streamlit

st.title("Gerador de Tabelas em Word")
st.write("Este aplicativo gera um arquivo Word com 7 tabelas específicas, de acordo com o banco de dados.")

# Botão para iniciar o processamento
if st.button("Gerar Relatório"):
    try:
        # Exibição de status usando Streamlit
        with st.spinner("Iniciando processo..."):
            st.info("Conectando ao banco de dados...", icon="🔌")

            # 1) Conectar ao banco e carregar dados
            conn = sqlite3.connect(CAMINHO_BD)
            df = pd.read_sql_query(f"SELECT * FROM {TABELA_BD}", conn)
            conn.close()

        st.success("Banco de dados carregado com sucesso!", icon="✅")

        # Barra de progresso
        progress_bar = st.progress(0, text="Processando dados...")

        # 2) Obter a data mais recente e filtrar últimos 5 meses
        # Converte a coluna 'data' em datetime (se estiver em formato yyyy-mm-dd)
        df['data_datetime'] = pd.to_datetime(df['data'], errors='coerce', format='%Y-%m-%d')
        data_maxima = df['data_datetime'].max()
        if pd.isnull(data_maxima):
            st.warning("Não foi encontrada data máxima válida no banco.", icon="⚠️")
            st.stop()

        cinco_meses_atras = data_maxima - relativedelta(months=5)
        df_filtrado = df[df['data_datetime'] >= cinco_meses_atras].copy()

        # 3) Criar um agrupamento por Tabela (nome da tabela final),
        #    usando o mapeamento placa -> tabela
        df_filtrado['nome_tabela_final'] = df_filtrado['placa'].map(placa_to_table)

        # 4) Preparar o documento do Word
        doc = Document()

        # Ajustar margens do documento
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

        # Nome do arquivo de saída (substitui se já existir)
        hoje = datetime.now().strftime("%d_%m_%y")
        nome_arquivo = f"tabelas_relatorio_{hoje}.docx"  # Removido acento
        nome_arquivo = sanitize_filename(nome_arquivo)

        # Verifica se diretório existe e cria se necessário
        if not os.path.exists(DIRETORIO_ARQUIVOS):
            os.makedirs(DIRETORIO_ARQUIVOS)
            st.info(f"Diretório para arquivos Word criado: {DIRETORIO_ARQUIVOS}")

        caminho_completo = os.path.join(DIRETORIO_ARQUIVOS, nome_arquivo)

        # Inicia a contagem para a barra de progresso
        total_tabelas = len(tabelas_relatorio)
        cont = 0

        for nome_tabela in tabelas_relatorio:
            cont += 1
            # Filtra somente as linhas daquela tabela
            df_tab = df_filtrado[df_filtrado['nome_tabela_final'] == nome_tabela].copy()

            if df_tab.empty:
                st.warning(f"Tabela '{nome_tabela}' está vazia (sem placas correspondentes ou sem dados recentes).", icon="⚠️")
            else:
                # Ordenar se desejar (por data ou por placa)
                df_tab.sort_values(by=["placa", "data_datetime"], inplace=True)

                # 5) Adiciona o título da tabela no Word
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_title = p.add_run(nome_tabela.upper())  # maiúsculo
                run_title.bold = True
                run_title.font.size = Pt(12)

                # Espaço entre o título e a tabela
                doc.add_paragraph("")

                # 6) Criar a tabela no Word
                #    -> 1 linha de cabeçalho + número de linhas = len(df_tab)
                #    -> 5 colunas: [Placa, Data, Cota, Coordenada Este, Coordenada Norte]
                table = doc.add_table(rows=1, cols=5)
                table.allow_autofit = True
                table.autofit = True

                # Configura cabeçalho
                hdr_cells = table.rows[0].cells
                hdr_texts = ["PLACA", "DATA", "COTA", "COORDENADA ESTE", "COORDENADA NORTE"]
                for i, txt in enumerate(hdr_texts):
                    # Texto no cabeçalho
                    paragraph = hdr_cells[i].paragraphs[0]
                    paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                    run = paragraph.add_run(txt)
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # branco

                    # Fundo azul no cabeçalho
                    tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), '0000FF')  # azul forte
                    tcPr.append(shd)

                # Força repetição do cabeçalho em múltiplas páginas
                set_repeat_table_header(table.rows[0])

                # 7) Popular dados na tabela
                for idx, row in df_tab.iterrows():
                    # Adiciona nova linha
                    row_cells = table.add_row().cells

                    # Formata placa
                    placa_str = str(row['placa']) if row['placa'] else ""

                    # Formata data
                    data_str = formatar_data_pt_br_iso(str(row['data']))

                    # Formata cota (elevação)
                    cota_str = ""
                    if row['elevacao'] is not None:
                        cota_str = formatar_cota(row['elevacao'])

                    # Formata coordenadas
                    este_str = ""
                    if row['coordenada_este'] is not None:
                        este_str = formatar_coordenada(row['coordenada_este'])

                    norte_str = ""
                    if row['coordenada_norte'] is not None:
                        norte_str = formatar_coordenada(row['coordenada_norte'])

                    row_cells[0].text = placa_str
                    row_cells[1].text = data_str
                    row_cells[2].text = cota_str
                    row_cells[3].text = este_str
                    row_cells[4].text = norte_str

                    # Ajustando fonte inicial (caso precise reduzir, o Word fará o autofit)
                    for c in row_cells:
                        for paragraph in c.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(11)  # Fonte inicial 11

                # 8) Quebra de página após cada tabela
                doc.add_page_break()

            # Atualiza barra de progresso
            progresso = int((cont / total_tabelas) * 100)
            progress_bar.progress(progresso, text=f"Gerando tabela {cont}/{total_tabelas}...")

        # 9) Salvar o arquivo (substituir se já existe)
        try:
            # Primeira tentativa - nome original
            doc.save(caminho_completo)
        except PermissionError:
            # Se o arquivo estiver em uso, tenta com timestamp adicional
            st.warning("Arquivo original está em uso. Tentando nome alternativo...", icon="⚠️")
            timestamp = datetime.now().strftime("%H%M%S")
            nome_alternativo = f"tabelas_relatorio_{hoje}_{timestamp}.docx"
            nome_alternativo = sanitize_filename(nome_alternativo)
            caminho_alternativo = os.path.join(DIRETORIO_ARQUIVOS, nome_alternativo)
            
            try:
                doc.save(caminho_alternativo)
                caminho_completo = caminho_alternativo  # Atualiza caminho para a mensagem final
                st.info(f"Arquivo salvo com nome alternativo: {nome_alternativo}")
            except Exception as e:
                st.error(f"Não foi possível salvar o arquivo mesmo com nome alternativo: {str(e)}", icon="❌")
                st.info("Sugestão: Verifique se você tem permissões de escrita na pasta ou se o arquivo está em uso.")
                raise

        # Mensagem de finalização
        st.success(f"Arquivo gerado com sucesso: {caminho_completo}", icon="✅")

    except Exception as e:
        st.error("Ocorreu um erro ao gerar o relatório.", icon="❌")
        st.exception(e)
