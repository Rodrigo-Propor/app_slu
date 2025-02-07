import sqlite3
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from datetime import datetime, timedelta
import locale

# Configurar locale para português do Brasil
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

# Definir cores e parâmetros globais
HEADER_BG_COLOR = "2F5496"  # Azul escuro para cabeçalho
BORDER_COLOR = "44B3E1"     # Azul claro para bordas
ALT_ROW_COLOR = "C0E6F5"    # Azul claro para linhas alternadas

def get_pampulha_filename(placa):
    """
    Determina o nome do arquivo baseado no padrão da placa para o local PAMPULHA
    """
    if not placa.startswith('PR '):
        return None
        
    identifier = placa[3:4]
    filename_map = {
        '1': 'AC01.docx',
        '3': 'AC03.docx',
        '4': 'AC04.docx',
        '5': 'AC05.docx',
        'A': 'Placas Ampliação.docx'
    }
    return filename_map.get(identifier)

def set_cell_margins(cell, top=0, right=100, bottom=0, left=100):
    """Define as margens das células"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for side, margin in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(margin))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    Define as bordas para uma célula
    Exemplo de uso: set_cell_border(cell, top={"sz": 4, "color": "44B3E1", "val": "single"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def format_table(table):
    """Aplica formatação completa à tabela"""
    # Configurar alinhamento da tabela
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Formatar cabeçalho
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            # Configurar cor de fundo do cabeçalho
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG_COLOR}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Formatar texto do cabeçalho
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Definir margens da célula
            set_cell_margins(cell)

    # Formatar células de dados
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            # Cor de fundo alternada
            bg_color = ALT_ROW_COLOR if i % 2 == 0 else "FFFFFF"
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Formatar texto
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                run = paragraph.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            
            # Definir margens da célula
            set_cell_margins(cell)
            
            # Configurar alinhamento vertical
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Aplicar bordas em todas as células
    border_props = {
        "sz": 4,
        "val": "single",
        "color": BORDER_COLOR,
    }
    
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=border_props,
                bottom=border_props,
                left=border_props,
                right=border_props,
            )

def format_number(value, decimals=4):
    """Formata números com vírgula decimal e ponto de milhar"""
    try:
        num = float(value)
        formatted = f"{num:,.{decimals}f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        return formatted
    except:
        return value

# Conectar ao banco de dados
db_path = 'banco_dados.db'
conn = sqlite3.connect(db_path)

# Encontrar a data mais recente e calcular data inicial (8 meses atrás)
query_ultima_data = '''
SELECT MAX(dt_data) as ultima_data
FROM dados_placa_geral
'''
ultima_data = pd.read_sql(query_ultima_data, conn)['ultima_data'].iloc[0]
data_inicial = (pd.to_datetime(ultima_data) - pd.DateOffset(months=8)).strftime('%Y-%m-%d')

# Buscar os dados dos últimos 8 meses
query = f'''
SELECT dt_data AS "Data", 
       ct_cota AS "Cota", 
       cd_este AS "Este", 
       cd_norte AS "Norte",
       lc_local AS "Local",
       pl_placa AS "Placa"
FROM dados_placa_geral
WHERE dt_data BETWEEN '{data_inicial}' AND '{ultima_data}'
ORDER BY lc_local, dt_data
'''
df = pd.read_sql(query, conn)
conn.close()

# Formatar dados
df['Data'] = pd.to_datetime(df['Data']).dt.strftime('%d/%m/%Y')
df['Cota'] = df['Cota'].apply(lambda x: format_number(x, 4))
df['Este'] = df['Este'].apply(lambda x: format_number(x, 4))
df['Norte'] = df['Norte'].apply(lambda x: format_number(x, 4))

# Criar um documento Word para cada local
for local in df['Local'].unique():
    # Se for PAMPULHA, precisamos subdividir por tipo de placa
    if local == 'PAMPULHA':
        # Agrupar dados por padrão de placa
        df_local = df[df['Local'] == local]
        processed_plates = set()  # Conjunto para rastrear placas já processadas
        
        for _, row in df_local.iterrows():
            placa = row['Placa']
            filename = get_pampulha_filename(placa)
            
            # Se já processamos esta placa ou não tem um padrão válido, pular
            if placa in processed_plates or not filename:
                continue
                
            # Filtrar dados para placas com o mesmo padrão
            plate_pattern = placa[:4]  # 'PR ' mais o primeiro caractere identificador
            df_specific = df_local[df_local['Placa'].str.startswith(plate_pattern)]
            
            doc = Document()
            
            # Configurar margens da página
            sections = doc.sections
            for section in sections:
                section.top_margin = Pt(50)
                section.bottom_margin = Pt(50)
                section.left_margin = Pt(50)
                section.right_margin = Pt(50)

            # Título
            title = doc.add_paragraph(f"Monitoramento - {local}")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(14)
            title_run.font.bold = True

            # Criar tabela com dados filtrados
            df_table = df_specific[['Placa', 'Data', 'Cota', 'Este', 'Norte']]
            table = doc.add_table(rows=1, cols=len(df_table.columns))
            
            # Adicionar cabeçalho
            header_names = ['Placa', 'Data', 'Cota', 'Coordenada E', 'Coordenada N']
            for i, column_name in enumerate(header_names):
                table.cell(0, i).text = column_name

            # Adicionar dados
            for _, data_row in df_table.iterrows():
                cells = table.add_row().cells
                for i, value in enumerate(data_row):
                    cells[i].text = str(value)

            # Aplicar formatação à tabela
            format_table(table)

            # Salvar documento
            doc.save(filename)
            print(f"Arquivo '{filename}' criado com sucesso.")
            
            # Marcar todas as placas deste padrão como processadas
            processed_plates.update(df_specific['Placa'].unique())
            
    else:
        # Processo normal para outros locais
        doc = Document()
        
        # Configurar margens da página
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(50)
            section.bottom_margin = Pt(50)
            section.left_margin = Pt(50)
            section.right_margin = Pt(50)

        # Título
        title = doc.add_paragraph(f"Monitoramento - {local}")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(14)
        title_run.font.bold = True

        # Filtrar dados do local e reordenar colunas
        df_local = df[df['Local'] == local][['Placa', 'Data', 'Cota', 'Este', 'Norte']]
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=len(df_local.columns))
        
        # Adicionar cabeçalho com nomes personalizados
        header_names = ['Placa', 'Data', 'Cota', 'Coordenada E', 'Coordenada N']
        for i, column_name in enumerate(header_names):
            table.cell(0, i).text = column_name

        # Adicionar dados
        for _, row in df_local.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

        # Aplicar formatação à tabela
        format_table(table)

        # Salvar documento
        doc.save(f"Monitoramento_{local.replace(' ', '_')}.docx")
        print(f"Arquivo 'Monitoramento_{local.replace(' ', '_')}.docx' criado com sucesso.")