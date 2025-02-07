from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from pathlib import Path
import os

# Definir cores e parâmetros globais
HEADER_BG_COLOR = "2F5496"  # Azul escuro para cabeçalho
BORDER_COLOR = "44B3E1"     # Azul claro para bordas
ALT_ROW_COLOR = "C0E6F5"    # Azul claro para linhas alternadas

def format_table(table):
    """
    Formata uma tabela Word com configurações específicas.
    """
    try:
        # Configura larguras das colunas
        adjust_column_widths(table)
        
        # Configurar propriedades da tabela
        set_table_properties(table)
        
        # Formatar linhas e células
        format_rows(table)
        
        # Aplicar bordas (removendo linhas internas)
        apply_borders(table)
        
        print("Formatação da tabela concluída com sucesso.")
    except Exception as e:
        print(f"Erro ao formatar tabela: {e}")

def adjust_column_widths(table):
    """
    Ajusta as larguras das colunas com base no conteúdo.
    """
    try:
        # Define proporções ideais para cada coluna
        width_ratios = [0.15, 0.15, 0.15, 0.275, 0.275]  # Total = 1
        total_width = sum(Inches(8.5).twips for _ in range(len(width_ratios)))
        
        for col_idx, col in enumerate(table.columns):
            # Calcula a largura ideal para a coluna
            width = int(total_width * width_ratios[col_idx])
            
            for cell in col.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(width))
                tcW.set(qn('w:type'), 'dxa')
                
                # Remove largura existente, se houver
                existing_tcW = tcPr.find(qn('w:tcW'))
                if existing_tcW is not None:
                    tcPr.remove(existing_tcW)
                
                tcPr.append(tcW)
    except Exception as e:
        print(f"Aviso: Erro ao ajustar larguras das colunas: {e}")

def remove_cell_borders(cell):
    """
    Remove todas as bordas de uma célula.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)

def set_table_properties(table):
    """
    Configura propriedades básicas da tabela.
    """
    try:
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Configuração do layout
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        existing_layout = tblPr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tblPr.remove(existing_layout)
        tblPr.append(tblLayout)
        
        # Centralização da tabela
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configuração do cabeçalho repetitivo
        if table.rows:
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            
            # Adiciona propriedade de cabeçalho
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            existing_header = trPr.find(qn('w:tblHeader'))
            if existing_header is not None:
                trPr.remove(existing_header)
            trPr.append(tblHeader)
            
            # Impede divisão da linha
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), "true")
            existing_cantSplit = trPr.find(qn('w:cantSplit'))
            if existing_cantSplit is not None:
                trPr.remove(existing_cantSplit)
            trPr.append(cantSplit)
    except Exception as e:
        print(f"Aviso: Erro ao configurar propriedades da tabela: {e}")

def format_rows(table):
    """
    Formata todas as linhas da tabela.
    """
    try:
        # Processa linha do cabeçalho
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # Remove bordas da célula
                remove_cell_borders(cell)
                
                # Aplica formatação do cabeçalho
                tcPr = cell._tc.get_or_add_tcPr()
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG_COLOR}"/>')
                existing_shd = tcPr.find(qn('w:shd'))
                if existing_shd is not None:
                    tcPr.remove(existing_shd)
                tcPr.append(shading_elm)
                
                # Formata texto do cabeçalho
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.bold = True
        
        # Processa demais linhas
        for i, row in enumerate(table.rows[1:], 1):
            bg_color = ALT_ROW_COLOR if i % 2 == 0 else "FFFFFF"
            for cell in row.cells:
                # Remove bordas da célula
                remove_cell_borders(cell)
                format_cell(cell, bg_color)
    except Exception as e:
        print(f"Aviso: Erro ao formatar linhas: {e}")

def format_cell(cell, bg_color):
    """
    Formata uma célula individual.
    """
    try:
        # Configuração do fundo
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        existing_shd = tcPr.find(qn('w:shd'))
        if existing_shd is not None:
            tcPr.remove(existing_shd)
        tcPr.append(shading_elm)
        
        # Configuração vertical
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Configuração do texto
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    except Exception as e:
        print(f"Aviso: Erro ao formatar célula: {e}")

def apply_borders(table):
    """
    Aplica bordas externas e remove bordas internas.
    """
    try:
        # Remove todas as bordas existentes primeiro
        for row in table.rows:
            for cell in row.cells:
                remove_cell_borders(cell)
        
        # Aplica apenas as bordas externas
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        
        borders = OxmlElement('w:tblBorders')
        
        # Define as bordas externas
        for border_type in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), BORDER_COLOR)
            borders.append(border)
        
        # Define explicitamente que não deve haver bordas internas
        for border_type in ['insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'none')
            borders.append(border)
        
        # Remove bordas existentes e adiciona as novas
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(borders)
    except Exception as e:
        print(f"Aviso: Erro ao aplicar bordas: {e}")

def process_directory():
    """
    Processa todos os arquivos .docx na pasta 'temporario'.
    """
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        temp_dir = os.path.join(script_dir, 'temporario')
        
        if not os.path.exists(temp_dir):
            print(f"Erro: A pasta 'temporario' não foi encontrada em {script_dir}")
            return
        
        docx_files = [f for f in Path(temp_dir).glob('*.docx') if not f.name.startswith('~$')]
        
        if not docx_files:
            print("Nenhum arquivo .docx encontrado na pasta 'temporario'")
            return
        
        print(f"Processando {len(docx_files)} arquivo(s)...")
        
        for file_path in docx_files:
            try:
                doc = Document(file_path)
                for table in doc.tables:
                    format_table(table)
                doc.save(file_path)
                print(f"Arquivo processado com sucesso: {file_path.name}")
            except Exception as e:
                print(f"Erro ao processar {file_path.name}: {e}")
        
        print("\nProcessamento concluído!")
    except Exception as e:
        print(f"Erro durante o processamento: {e}")

if __name__ == "__main__":
    print("Formatador de Tabelas em Documentos Word")
    print("=======================================")
    process_directory()
    input("\nPressione Enter para sair...")